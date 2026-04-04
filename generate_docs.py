# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_presentation_script():
    doc = Document()
    
    # Title
    title = doc.add_heading('Kịch Bản Thuyết Trình Bảo Vệ Đồ Án Data Science', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Dự án: Dự đoán giá & phát hiện bất thường giá nhà ở (Nền tảng Nhà Tốt)\nNhóm: DL07 - K311 (Đoàn Nhật Quang – Phan Ngọc Minh Quân)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_slide(num, content, speaker="Quang & Quân"):
        h = doc.add_heading(f'Slide {num}', level=1)
        p = doc.add_paragraph()
        run = p.add_run(f'[{speaker}]: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)
        p.add_run(content)

    content_list = [
        (1, "Dạ em chào Cô và các bạn. Tụi em là nhóm số 9, gồm 2 thành viên là Nhật Quang và Minh Quân. Hôm nay tụi em xin phép được trình bày đồ án tốt nghiệp Data Science với đề tài: 'Dự đoán giá và phát hiện bất thường giá nhà ở trên nền tảng Nhà Tốt - khu vực HCM'.", "Quang"),
        (2, "Mục lục của bài thuyết trình hôm nay sẽ đi qua 7 phần chính. 6 phần theo chuẩn quy trình quy trình CRISP-DM từ lúc hiểu bài toán, thu thập, làm sạch dữ liệu, EDA, Modeling, và phát hiện bất thường. Đặc biệt có một phần Bonus nho nhỏ ở phần 7 về cách nhóm tự cấu trúc hệ thống Crawl dữ liệu quy mô lớn.", "Quang"),
        (3, "Đầu tiên, em xin phép đi vào phần Business Objective - Bối cảnh và Mục tiêu của dự án.", "Quang"),
        (4, "Hiện tại Nhà Tốt đang là nền tảng BĐS rất lớn thuộc Chợ Tốt. Vấn đề thực tế đặt ra là: Người bán thì luôn băn khoăn không biết định giá nhà mình bao nhiêu cho hợp lý để mau trôi; trong khi Nền tảng thì phải vất vả duyệt tin, lọc các tin rao giá ảo dìm thị trường hoặc lừa đảo.\n\nTừ đó, nhóm em xác định 2 bài toán lớn. Bài toán 1: Xây mô hình dự đoán giá bán nhà để hỗ trợ người bán. Bài toán 2: Phân tích độ lệch chuẩn để cảnh báo các bài đăng có mức giá bất thường giúp nền tảng kiểm duyệt tự động.", "Quang"),
        (5, "Để giải quyết 2 bài toán đó, nhóm bước vào phần Thu thập và Làm sạch dữ liệu.", "Quân"),
        (6, "Bên tay trái là giao diện của nền tảng Nhà Tốt. Cột bên tay phải là các trường dữ liệu thô mà mình có thể trích xuất được. Nhóm tập trung vào các trường quan trọng mang tính quyết định đến giá như: Diện tích, Số phòng ngủ, Loại hình, Vị trí, Tính pháp lý.", "Quân"),
        (7, "Đây là cái phễu (pipeline) xuyên suốt của nhóm. Bắt đầu từ 8,273 mẫu tin gốc rải đều ở 3 quận Bình Thạnh, Gò Vấp, Phú Nhuận do Gò Vấp chiếm tận 53%. Qua quá trình làm sạch drop missing value > 30%, quy mô gọn lại còn gần 8,000 dòng. Đến bước Feature Engineering, nhóm biến đổi và nội suy tính toán từ 9 cột ra thành tận 29 features rất giá trị.", "Quân"),
        (8, "Data frame hiển thị ra bảng trông như thế này.", "Quân"),
        (9, "Tiếp tục slide hiển thị bảng dữ liệu Data frame.", "Quân"),
        (10, "Tiếp theo, nhóm tiến hành EDA (Exploratory Data Analysis) và Feature Engineering.", "Quang"),
        (11, "Khi phân tích, nhóm thấy giá bán trung vị ở mức 6.69 tỷ nhưng lại lệch phải rất mạnh (tức là có quá nhiều nhà siêu đắt kéo lệch phân phối). Do đó, nhóm bắt buộc phải lấy Logarit tự nhiên của Giá ở các bước Modeling. Nhóm cũng tìm được các đặc trưng tương quan mạnh với giá như diện tích (+0.73), số phòng ngủ, và đặc biệt là nhà có thang máy kéo giá trị lên rất cao.", "Quang"),
        (12, "Đi sâu vào từng quận thì Phú Nhuận đắt nhất, Gò Vấp rẻ nhất. Phân khúc 4 đến 7 tỷ đồng chiếm sóng chính. Nếu nhà ở mặt tiền, giá tự động vọt lên, chưa kể nếu có thang máy, giá nhà bị đẩy thêm trung bình 7.5 tỷ.", "Quang"),
        (13, "Từ lúc phân tích sâu như vậy, nhóm thực hiện kỹ thuật Feature Engineering cực kỳ chi tiết. Nhóm sinh được 29 tính năng. Đặc biệt phải kể đến 16 tính năng text - nhóm dùng Regex quét thẳng vào Cột Tiêu Đề và Mô Tả để rút trích các cụm từ vàng như 'thang máy', 'sân thượng', 'hẻm xe hơi'. Đây là các insight mà nếu chỉ nhìn features số sẽ rớt mất.", "Quang"),
        (14, "Minh họa bảng tính năng sau khi sinh thêm các trường bool qua Regex.", "Quang"),
        (15, "Minh họa bảng tính năng sinh bằng Regex (tiếp theo).", "Quang"),
        (16, "Cách quét Regex để lọc ra nhà lô góc, đang cho thuê, giáp mặt tiền.", "Quang"),
        (17, "Sử dụng Regex dò tìm nhà nở hậu, kinh doanh đỉnh, và yếu tố 'có thang máy'.", "Quang"),
        (18, "Lọc thêm các thông tin rủi ro/khẩn cấp như 'ngộp ngân hàng', 'bán gấp' hoặc 'nhà mới ở ngay'.", "Quang"),
        (19, "Đây là bảng tóm tắt 16 binary features sau khi áp Regex, kèm theo đó là điểm tiện nghi (tiện nghi score) được tích lũy lại.", "Quang"),
        (20, "Các bảng đồ thị phân phối giá sau khi Logarit.", "Quang"),
        (21, "Và giờ là phần cốt lõi: Modeling dự đoán giá. Nhóm chạy đồng thời trên 2 môi trường: thư viện Sklearn kiểu truyền thống và công nghệ dữ liệu lớn PySpark MLlib.", "Quân"),
        (22, "Trên Sklearn, nhóm chạy 4 phương pháp. Kết quả rực rỡ nhất gọi tên Random Forest với độ chính xác theo dạng R² đạt 84.6%, sai lệch giá chỉ rơi vào quanh mốc 14%. Nếu chuyển ra tiền mặt trên thực tế thì MAE lệch loanh quanh 1 tỷ đồng. Diện tích và Số phòng ngủ là 2 yếu tố học được lớn nhất.", "Quân"),
        (23, "Đồ thị Scatter: Giá dự đoán bám cực kỳ sát Giá thực tế chạy chéo góc 45 độ.", "Quân"),
        (24, "Đồ thị so sánh các Model trên Sklearn. Random Forest giữ vị trí vô địch ổn định.", "Quân"),
        (25, "Chi tiết sai số học tập và validation của Random Forest trên Sklearn.", "Quân"),
        (26, "Còn bên vũ trụ Big Data PySpark? Cũng 4 model trên framework phân tán. Random Forest vẫn chiến thắng với R² đạt 82%. MAE lúc này nhỉnh hơn một chút (1.3 tỷ) với thời gian train cực kỳ nhanh so với Gradient Boosting Tree (59 giây vs 315 giây).", "Quân"),
        (27, "Kết luận cho phần dự đoán giá ở cả 2 môi trường: Random Forest kháng quá khớp (overfitting) cực tốt, tốc độ và độ cân bằng vượt trội, xứng đáng làm model chạy chính cho production server của hệ thống Nhà Tốt.", "Quân"),
        (28, "Tóm lại, nếu dữ liệu dưới 10k dòng như hiện tại, Sklearn đem lại kết quả tốt hơn, dễ config thuật toán chuyên sâu để ăn điểm. Nhưng nếu hệ thống đẩy lên tới vài chục triệu tin, PySpark sẽ là vị vua không thể thay thế vì tính mở rộng (Scalable) cluster.", "Quang"),
        (29, "Tiếp nối dự đoán giá, chúng ta sang Bài toán số 2: Phát hiện tin đăng bất thường (Anomaly Detection).", "Quang"),
        (30, "Nguyên lý phát hiện của tụi em không dựa vào 1 thuật toán cảm tính rẻ tiền. Tụi em xây hẳn 1 công thức 'Composite Score' (điểm tổng hợp) chia tỷ trọng dựa trên 4 trụ cột. 40% soi độ lệch Resid Z-score. 10% vi phạm viền giá trần-sàn thực tế. 20% nằm ngoài mốc Percentile 10-90. Và 30% chốt chặn bằng công nghệ cách ly bất thường đa chiều Machine Learning Unsupervised là Isolation Forest. Cứ ai trên 70 điểm là gán nhãn Đỏ (Red - Cực kỳ bất thường).", "Quang"),
        (31, "Cùng nhìn lại những 'thủ phạm' bị bắt quả tang trên Sklearn: Điểm cao chót vót. Có một căn nhà ở Gò Vấp 106m2 mà giá thực là tự dưng la làng bán 142 Tỷ. Hệ thống nhóm em predict ra giá trị cao cực đại của nó chỉ được 53 tỷ, thế là điểm báo độc hại > 81đ. Hay 1 nhà nhỏ 18m2 Gò Vấp rẻ mạt dưới 1 tỷ (quá thấp).", "Quang"),
        (32, "Data view Top 10 tin bất thường trên Sklearn.", "Quang"),
        (33, "Đồ thị 3D thể hiện mật độ rải rác của các điểm bất thường Sklearn.", "Quang"),
        (34, "Đồ thị chi tiết Anomalies Sklearn đánh dấu điểm đỏ outliers.", "Quang"),
        (35, "Tương tự, hiển thị Top 10 tin bất thường được bắt bởi PySpark.", "Quang"),
        (36, "Thống kê PySpark cho thấy, 91.2% tin là bình thường (sạch), có 44 tin chiếm 0.6% là Rất Bất Thường trên bảng cảnh báo. Đa phần đánh rớt vì lý do tự ý đôn giá CAO vượt thực tế.", "Quang"),
        (37, "Biểu đồ Bar chart trên PySpark. Số lượng tin cao ảo vọt lên rất rõ.", "Quang"),
        (38, "Scatter phân phối Anomalies trên nền tảng PySpark.", "Quang"),
        (39, "Dạ vừa rồi là 2 bài toán lớn, em xin đi đến Kết luận và Hướng triển khai thực tế.", "Quân"),
        (40, "Tựu chung lại, đồ án cung cấp một Data pipeline end-to-end mạch lạc. Thành quả là bộ máy Random Forest R^2 > 0.82 và Hệ thống quy chuẩn Composite Score cô lập tin bất hợp lý.", "Quân"),
        (41, "Từ bài toán này, thực tế Nhà Tốt hay Chợ Tốt có thể áp dụng lập tức. Cho người bán: Nền tảng tự sinh Suggestion Box - Giá Gợi Ý ngay khi dân điền khung nhập chữ. Cho bản thân Website, có ngay Filter lọc tin scam tự động để ẩn bài ẩn tương tác của người đăng.", "Quân"),
        (42, "Thay vì phụ thuộc data có sẵn, nhóm em muốn trình bày 1 phần Bonus cực kỳ thú vị: Cách Nhóm tự lập trình Crawler lấy data Live trực tiếp.", "Quang"),
        (43, "Ngày nay mấy trang lớn cấu hình tường lửa Anti-bot 403, Cloudflare chống cào dữ liệu rát dữ dội. Nhưng team em sử dụng kỹ thuật giả mạo dấu vân tay mã hóa TLS Fingerprint bằng cffi giả lập Chrome bản quyền, cộng thêm Session Warmup đi cửa chính thì bypass vượt mọi rào cản thành công.", "Quang"),
        (44, "Cứ thế tụi em vòng lặp Web Scrape kết hợp API Gateway dự phòng, bắt dữ liệu từ 24 quận huyện TP HCM, với gần 60.000 records. Nguồn data này có thể mở rộng toàn diện sau này.", "Quang"),
        (45, "Source code Crawler viết bằng Python nằm trong repo của dự án.", "Quang"),
        (46, "Tiếp theo là phần Phân công công việc. Xin lỗi giảng viên vì là 2 người tự cày nguyên end-to-end framework nên bọn em phải kiêm nghiệm, thay nhau double check từng file từ Crawl cho đến Modeling nên hầu hết là cả 2 cùng nắm chi tiết.", "Quân"),
        (47, "Sân khấu thành viên nhóm: Đoàn Nhật Quang & Phan Ngọc Minh Quân.", "Quân"),
        (48, "Lý do chọn khóa học là nhóm không chỉ muốn xây back-end system chay mà còn đẩy kiến thức Data Science vào nhân lõi hệ sinh thái.", "Quang"),
        (49, "Kinh nghiệm to bự nhất được rút ra sau môn này là cách setup Data Pipeline luân chuyển dữ liệu và tư duy MLOps hiện đại.", "Quân"),
        (50, "Đồ án rất tâm huyết, phần trình bày của nhóm 9 đến đây là kết thúc. Dạ chúng em rất mong nhận được những góp ý, câu hỏi của Cô và các anh chị lớp mình! CẢM ƠN cô và các bạn!", "Quang & Quân")
    ]
    
    for num, content, speaker in content_list:
        add_slide(num, content, speaker)
        
    doc.save('presentation.docx')
    print("Done generating presentation.docx")

def create_qa_script():
    doc = Document()
    title = doc.add_heading('Bộ Câu Hỏi Hỏi Đáp (Q&A) - Bảo Vệ Đồ Án', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    questions = [
        ("Tại sao nhóm chọn lấy Log(giá_bán) làm Target dự đoán thay vì để nguyên giá trị thô, cuối cùng khi hiển thị ra cho khách hàng nhóm xử lý như thế nào?",
         "Việc dùng giá thô (raw price) gặp vấn đề rất lớn vì phân phối giá BĐS bị lệch chuẩn cực nặng về bên phải (Right Skewed, skewness = 6.49). Điều này khiến model MSE bị nhạy cảm quá mức với các căn siêu biệt thự (vài trăm tỷ) và bỏ rơi tính ưu việt của phân khúc nhà giá rẻ.\nViệc biến đổi sang Logarit giúp nén và kéo phân phối về hình chuông chuẩn hơn (normal distribution, skew=0.26). Nhờ đó, Linear regression và Ensemble Trees học rất mượt. Khi show cho user (hiển thị giao diện), hệ thống chỉ việc sử dụng hàm mũ np.exp(y_predict) để khôi phục tiền tỷ ra giá trị thực."),
        
        ("Nhóm trích xuất (Feature Engineering) ra 16 binary text features từ `mô tả` và `tiêu đề` thông qua Regex, liệu cách này có bị bắt nhầm ví dụ: câu 'không có thang máy' mà regex báo 'có thang máy' không?",
         "Dạ một câu hỏi rất thực tế! Trong bước Regex, nhóm xử lý cẩn thận bằng cách sử dụng look-around assertion. Để lọc chữ 'thang máy', biểu thức Regex giới hạn việc không bắt qua các khoảng đứt câu có từ phủ định như 'không', 'chưa được'. Tuy nhiên, regex quy mô lớn luôn có tỷ lệ sai số nho nhỏ. May mắn thay, mô hình Tree-based model (Random Forest) có khả năng cô lập được noise cục bộ, nên qua thực nghiệm, 16 features do Regex rà được đã gánh accuracy của toàn mô hình tăng thêm tầm 5% so với không dùng."),
        
        ("Tại sao trên PySpark, kết quả R² của các model lại vô tình có vẻ thua sút hơn Sklearn (82% vs 84.6%)? Mặc dù PySpark chuyên trị quy mô đại trà?",
         "Dạ, PySpark cực kỳ mạnh về khả năng chia để trị (Distributed Computing) đối với Big Data với nhiều Workers. Nhưng bản thân dataset đồ án nhóm mình vì đã lọc missing value nên dữ liệu rơi rớt về 6,979 dòng. Bộ dữ liệu này nằm vừa vặn vào RAM của một máy đơn (Local processing). Việc này khiến Sklearn phát huy toàn cảnh bức tranh (global view) và tinh chỉnh siêu tham số mượt hơn. Đồng thời, thư viện MLlib trong PySpark không cài native XGBoost hay LightGBM tinh vi như Scikit-learn mà đa số là implementation gốc, nên kết quả mô hình nhỏ trong PySpark có độ R2 lui lại 1 bước nhưng bù lại khả năng scale-up vô cực."),
        
        ("Trong bài toán Anomaly Detection, nhóm đưa ra 1 cái Composite Score gồm 4 thành phần (Residual Z-Score, MinMax, Percentile, và Isolaton Forest). Căn cứ theo tỷ trọng 40% - 10% - 20% - 30%, việc này là có tính khoa học hay do cảm tính?",
         "Dạ cơ chế đánh trọng số này được Calibration thông qua thử nghiệm phân tích nhân tố chủ quan. Nghĩa là nhóm xem xét độ nhạy của tín hiệu rủi ro.\n- 40% Residual Z-Score: Do Random Forest có khả năng nội suy quá tốt, độ lệch trên phần dư này phản ánh sự sai lệch cung cầu thị trường mãnh liệt.\n- 30% Isolation Forest: Yếu tố Unsupervised ML. Bắt các outliners đa chiều mà rules cứng khó lòng phát giác.\n- 30% còn lại là các rules vi phạm toán học cứng (vi phạm Min/Max và Percentile).\nViệc blend theo 4-1-2-3 giúp tín hiệu còi hú báo bất thường ổn định, không bị chém nhầm khi giá nhà đắt chỉ đơn giản do nó xịn."),
         
        ("Crawling Data bằng kĩ thuật TLS Fingerprint (curl_cffi) chống anti-bot hoạt động ra sao? Khác cơ chế crawl thường như Selenium thế nào?",
         "Dạ nếu dùng Selenium hay Requests thuần túy, trình phân giải Cloudflare của bên Nhà Tốt sẽ gạt IP của mình ra khỏi traffic lập tức vì Header và mã bắt tay SSL/TLS Certificate quá sơ sài và đậm mùi 'lập trình bot'.\nNhóm em dùng `curl_cffi` giúp giả mạo toàn bộ hệ giá trị chuỗi byte handshake giống hệt như một thẻ tab từ trình duyệt Chrome 131 của người thật mở ra. Do đó, Bypass sạch bot-protection mà không cần load Headless Browser cồng kềnh ngốn RAM như Selenium. Việc chạy theo Session tự động exponential back-offs khi bị 429 Error khiến việc Crawl ngầm hàng triệu records diễn ra bí mật và trơn tru."),
         
        ("Với MAE của Random Forest là 1.07 tỷ, nhóm đánh giá kết quả MAE này là tệ hay chấp nhận được cho Bài toán định giá nhà?",
         "Dạ sai lệch khoảng trung bình 1 tỷ nghe khá lớn, nhưng so trên hệ quy chiếu giá nhà TPHCM phân khúc 6-15 tỷ thì độ lệch dải 1 tỷ tương đương cỡ 10% - 15%. Đây là mức dao động có thể chấp nhận được do tính thanh khoản và đàm phán trong Bất Động Sản. Giá đăng bán không phải là giá giao dịch thật sự. Nền tảng chỉ cần khuyến cáo 'Nhà của bạn có khả năng bán ở mốc 7 tỷ +- 500tr'. Điều này vẫn cực kỳ có lợi cho việc định hướng thị trường.")
    ]
    
    for idx, (q, a) in enumerate(questions, 1):
        qp = doc.add_paragraph()
        q_run = qp.add_run(f"Câu hỏi {idx}: {q}")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(192, 0, 0)
        
        ap = doc.add_paragraph()
        a_run = ap.add_run(f"Trả lời: {a}")
        ap.paragraph_format.left_indent = Pt(20)
        
    doc.save('question.docx')
    print("Done generating question.docx")

if __name__ == '__main__':
    create_presentation_script()
    create_qa_script()
